"""Word document generator for interview preparation materials."""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """Generator for creating Word documents with interview preparation content."""

    def __init__(self) -> None:
        """Initialize the Word document generator."""
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)

    def generate_interview_prep_document(
        self,
        prep_content: Dict[str, Any],
        job_title: str,
        company_name: Optional[str] = None,
    ) -> Optional[str]:
        """
        Generate a comprehensive interview preparation Word document.
        
        Args:
            prep_content: Interview preparation content dictionary
            job_title: Job title for the position
            company_name: Company name if available
            
        Returns:
            Path to generated document or None if generation fails
        """
        try:
            # Create new document
            doc = Document()
            
            # Set up styles
            self._setup_document_styles(doc)
            
            # Add header
            self._add_document_header(doc, job_title, company_name)
            
            # Add table of contents
            self._add_table_of_contents(doc)
            
            # Add company analysis section
            self._add_company_analysis_section(doc, prep_content.get("company_analysis", {}))
            
            # Add interview questions section
            self._add_interview_questions_section(doc, prep_content.get("interview_questions", []))
            
            # Add answer frameworks section
            self._add_answer_frameworks_section(doc, prep_content.get("suggested_answers", {}))
            
            # Add STAR stories section
            self._add_star_stories_section(doc, prep_content.get("star_stories", []))
            
            # Add questions to ask section
            self._add_questions_to_ask_section(doc, prep_content.get("questions_to_ask", []))
            
            # Add salary insights section
            self._add_salary_insights_section(doc, prep_content.get("salary_insights", {}))
            
            # Add overqualification tips section
            if prep_content.get("overqualification_tips"):
                self._add_overqualification_section(doc, prep_content["overqualification_tips"])
            
            # Add footer
            self._add_document_footer(doc)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Limit job title to first 50 characters and clean it
            truncated_title = job_title[:50] if job_title else "Job"
            safe_job_title = "".join(c for c in truncated_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_company = ""
            if company_name:
                truncated_company = company_name[:30]  # Limit company name too
                safe_company = f"_{truncated_company}"
                safe_company = "".join(c for c in safe_company if c.isalnum() or c in (' ', '-', '_')).rstrip()
            
            # Ensure total filename length is reasonable (max 150 chars)
            base_name = f"Interview_Prep_{safe_job_title}{safe_company}_{timestamp}"
            if len(base_name) > 140:  # Leave room for .docx
                safe_job_title = safe_job_title[:30]
                base_name = f"Interview_Prep_{safe_job_title}{safe_company}_{timestamp}"
            
            filename = f"{base_name}.docx"
            filepath = self.output_dir / filename
            
            doc.save(str(filepath))
            logger.info(f"Interview preparation document saved: {filepath}")
            
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            return None

    def generate_complete_analysis_document(
        self,
        analysis_content: Dict[str, Any],
        job_title: str,
        company_name: Optional[str] = None,
    ) -> Optional[str]:
        """
        Generate a comprehensive analysis document with all results.
        
        Args:
            analysis_content: Complete analysis content dictionary
            job_title: Job title for the position
            company_name: Company name if available
            
        Returns:
            Path to generated document or None if generation fails
        """
        try:
            # Create new document
            doc = Document()
            
            # Set up styles
            self._setup_document_styles(doc)
            
            # Add header
            self._add_analysis_header(doc, analysis_content, job_title, company_name)
            
            # Add analysis summary
            self._add_analysis_summary_section(doc, analysis_content)
            
            # Add score breakdown
            self._add_score_breakdown_section(doc, analysis_content)
            
            # Add strong points
            self._add_strong_points_section(doc, analysis_content)
            
            # Add weak points
            self._add_weak_points_section(doc, analysis_content)
            
            # Add recommendations
            self._add_recommendations_section(doc, analysis_content)
            
            # Add interview preparation
            self._add_interview_prep_section(doc, analysis_content)
            
            # Add footer
            self._add_document_footer(doc)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            truncated_title = job_title[:30] if job_title else "Analysis"
            safe_job_title = "".join(c for c in truncated_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_company = ""
            if company_name:
                truncated_company = company_name[:20]
                safe_company = f"_{truncated_company}"
                safe_company = "".join(c for c in safe_company if c.isalnum() or c in (' ', '-', '_')).rstrip()
            
            filename = f"CV_Analysis_{safe_job_title}{safe_company}_{timestamp}.docx"
            filepath = self.output_dir / filename
            
            doc.save(str(filepath))
            logger.info(f"Complete analysis document saved: {filepath}")
            
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generating complete analysis document: {str(e)}")
            return None

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        styles = doc.styles
        
        # Create custom heading styles
        try:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Inches(0.25)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Inches(0.2)
        except:
            pass  # Style might already exist
        
        try:
            section_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Arial'
            section_style.font.size = Inches(0.18)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 51, 102)
            section_style.paragraph_format.space_before = Inches(0.15)
            section_style.paragraph_format.space_after = Inches(0.1)
        except:
            pass

    def _add_document_header(self, doc: Document, job_title: str, company_name: Optional[str]) -> None:
        """Add document header with title and metadata."""
        # Main title
        title = f"Interview Preparation Guide"
        title_para = doc.add_paragraph(title)
        try:
            title_para.style = 'CustomTitle'
        except:
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].bold = True
            title_para.runs[0].font.size = Inches(0.25)
        
        # Subtitle
        subtitle = f"Position: {job_title}"
        if company_name:
            subtitle += f" at {company_name}"
        
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].italic = True
        
        # Date
        date_para = doc.add_paragraph(f"Prepared on: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents."""
        toc_title = doc.add_paragraph("Table of Contents")
        try:
            toc_title.style = 'CustomHeading'
        except:
            toc_title.runs[0].bold = True
            toc_title.runs[0].font.size = Inches(0.18)
        
        toc_items = [
            "1. Company and Role Analysis",
            "2. Predicted Interview Questions",
            "3. Answer Frameworks and Strategies",
            "4. STAR Stories Template",
            "5. Questions to Ask the Interviewer",
            "6. Salary Negotiation Insights",
            "7. Handling Overqualification Concerns",
        ]
        
        for item in toc_items:
            toc_para = doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()

    def _add_company_analysis_section(self, doc: Document, company_analysis: Dict[str, str]) -> None:
        """Add company and role analysis section."""
        section_title = doc.add_paragraph("1. Company and Role Analysis")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        if company_analysis:
            for key, value in company_analysis.items():
                if value:
                    key_formatted = key.replace("_", " ").title()
                    para = doc.add_paragraph()
                    para.add_run(f"{key_formatted}: ").bold = True
                    para.add_run(value)
        
        # Add research tips
        doc.add_paragraph("\nResearch Tips:", style='Heading 3')
        tips = [
            "Visit the company website and read recent news/press releases",
            "Check the company's LinkedIn page for recent updates and employee insights",
            "Look up the hiring manager or team members on LinkedIn",
            "Research the company's competitors and market position",
            "Understand the company's products, services, and target customers",
        ]
        
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')

    def _add_interview_questions_section(self, doc: Document, questions: list) -> None:
        """Add predicted interview questions section."""
        section_title = doc.add_paragraph("\n2. Predicted Interview Questions")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        if questions:
            current_category = ""
            for question in questions:
                category = question.get("category", "General")
                if category != current_category:
                    current_category = category
                    cat_title = doc.add_paragraph(f"\n{category} Questions:", style='Heading 3')
                
                q_para = doc.add_paragraph()
                q_para.add_run("Q: ").bold = True
                q_para.add_run(question.get("question", ""))
                
                # Add space after each question
                doc.add_paragraph("")

    def _add_answer_frameworks_section(self, doc: Document, suggested_answers: Dict[str, str]) -> None:
        """Add answer frameworks and strategies section."""
        section_title = doc.add_paragraph("\n3. Answer Frameworks and Strategies")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        if suggested_answers:
            for framework_name, framework_content in suggested_answers.items():
                framework_title = framework_name.replace("_", " ").title()
                doc.add_paragraph(f"\n{framework_title}:", style='Heading 3')
                
                # Split content into paragraphs for better formatting
                paragraphs = framework_content.strip().split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

    def _add_star_stories_section(self, doc: Document, star_stories: list) -> None:
        """Add STAR stories template section."""
        section_title = doc.add_paragraph("\n4. STAR Stories Template")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        doc.add_paragraph("Prepare specific examples using the STAR method (Situation, Task, Action, Result):")
        
        for story in star_stories:
            story_title = doc.add_paragraph(f"\n{story.get('title', 'Story')}:", style='Heading 3')
            
            for component in ['situation', 'task', 'action', 'result']:
                if component in story:
                    comp_para = doc.add_paragraph()
                    comp_para.add_run(f"{component.title()}: ").bold = True
                    comp_para.add_run(story[component])

    def _add_questions_to_ask_section(self, doc: Document, questions_to_ask: list) -> None:
        """Add questions to ask the interviewer section."""
        section_title = doc.add_paragraph("\n5. Questions to Ask the Interviewer")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        doc.add_paragraph("Asking thoughtful questions shows your interest and helps you evaluate the role:")
        
        current_category = ""
        for q_item in questions_to_ask:
            category = q_item.get("category", "General")
            if category != current_category:
                current_category = category
                doc.add_paragraph(f"\n{category}:", style='Heading 3')
            
            q_para = doc.add_paragraph()
            q_para.add_run("• ").bold = True
            q_para.add_run(q_item.get("question", ""))
            
            if q_item.get("purpose"):
                purpose_para = doc.add_paragraph(f"  Purpose: {q_item['purpose']}")
                purpose_para.runs[0].italic = True

    def _add_salary_insights_section(self, doc: Document, salary_insights: Dict[str, str]) -> None:
        """Add salary negotiation insights section."""
        section_title = doc.add_paragraph("\n6. Salary Negotiation Insights")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        if salary_insights:
            for insight_key, insight_value in salary_insights.items():
                insight_title = insight_key.replace("_", " ").title()
                insight_para = doc.add_paragraph()
                insight_para.add_run(f"{insight_title}: ").bold = True
                insight_para.add_run(insight_value)

    def _add_overqualification_section(self, doc: Document, overqualification_tips: list) -> None:
        """Add overqualification handling section."""
        section_title = doc.add_paragraph("\n7. Handling Overqualification Concerns")
        try:
            section_title.style = 'CustomHeading'
        except:
            section_title.runs[0].bold = True
        
        doc.add_paragraph("If your background might be seen as overqualification, prepare for these concerns:")
        
        for tip in overqualification_tips:
            concern_para = doc.add_paragraph()
            concern_para.add_run("Concern: ").bold = True
            concern_para.add_run(tip.get("concern", ""))
            
            strategy_para = doc.add_paragraph()
            strategy_para.add_run("Strategy: ").bold = True
            strategy_para.add_run(tip.get("strategy", ""))
            
            if tip.get("example"):
                example_para = doc.add_paragraph()
                example_para.add_run("Example response: ").bold = True
                example_para.add_run(f'"{tip["example"]}"')
                example_para.runs[-1].italic = True
            
            doc.add_paragraph("")  # Add spacing

    def _add_document_footer(self, doc: Document) -> None:
        """Add document footer with tips."""
        doc.add_page_break()
        
        footer_title = doc.add_paragraph("Final Tips for Success")
        try:
            footer_title.style = 'CustomHeading'
        except:
            footer_title.runs[0].bold = True
        
        final_tips = [
            "Practice your answers out loud before the interview",
            "Prepare specific examples that demonstrate your skills and achievements",
            "Research the interviewer(s) on LinkedIn if their names are provided",
            "Prepare questions that show your genuine interest in the role and company",
            "Plan your route and arrive 10-15 minutes early",
            "Bring multiple copies of your resume and a notepad",
            "Follow up with a thank-you email within 24 hours",
        ]
        
        for tip in final_tips:
            doc.add_paragraph(tip, style='List Bullet')
        
        # Add generation timestamp
        timestamp_para = doc.add_paragraph(f"\nDocument generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        timestamp_para.runs[0].italic = True

    def _add_analysis_header(self, doc: Document, analysis_content: Dict[str, Any], job_title: str, company_name: Optional[str]) -> None:
        """Add header for analysis document."""
        # Main title
        title = doc.add_paragraph("CV ANALYSIS REPORT")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            title.style = 'Title'
        except:
            title.runs[0].bold = True
            title.runs[0].font.size = Inches(0.2)
        
        # Subtitle
        summary = analysis_content.get("analysis_summary", {})
        subtitle = f"Analysis for {summary.get('candidate_name', 'Candidate')}"
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            subtitle_para.style = 'Subtitle'
        except:
            subtitle_para.runs[0].bold = True
        
        # Job details
        job_details = f"Position: {job_title[:100]}"
        if company_name:
            job_details += f" | Company: {company_name}"
        job_details += f" | Analysis Date: {summary.get('analysis_date', 'N/A')}"
        
        details_para = doc.add_paragraph(job_details)
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details_para.runs[0].italic = True

    def _add_analysis_summary_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add analysis summary section."""
        doc.add_paragraph("")
        
        summary_title = doc.add_paragraph("📊 COMPATIBILITY ANALYSIS SUMMARY")
        try:
            summary_title.style = 'CustomHeading'
        except:
            summary_title.runs[0].bold = True
        
        summary = analysis_content.get("analysis_summary", {})
        score = summary.get("score", 0)
        
        # Score display
        score_para = doc.add_paragraph()
        score_para.add_run("Compatibility Score: ").bold = True
        score_run = score_para.add_run(f"{score}/100")
        score_run.bold = True
        score_run.font.size = Inches(0.15)
        
        # Score interpretation
        if score >= 80:
            interpretation = "Excellent match! Your profile aligns very well with the requirements."
        elif score >= 60:
            interpretation = "Good match with some areas for improvement."
        elif score >= 40:
            interpretation = "Moderate match. Focus on addressing key gaps."
        else:
            interpretation = "Significant improvements needed to match requirements."
        
        doc.add_paragraph(f"Assessment: {interpretation}")

    def _add_score_breakdown_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add score breakdown section."""
        doc.add_paragraph("")
        
        breakdown_title = doc.add_paragraph("🔍 DETAILED SCORE BREAKDOWN")
        try:
            breakdown_title.style = 'CustomHeading'
        except:
            breakdown_title.runs[0].bold = True
        
        breakdown = analysis_content.get("score_breakdown", {})
        for category, details in breakdown.items():
            if isinstance(details, dict) and "score" in details:
                para = doc.add_paragraph()
                para.add_run(f"{category.replace('_', ' ').title()}: ").bold = True
                para.add_run(f"{details['score']}/100")
                if details.get("reason"):
                    para.add_run(f" - {details['reason']}")

    def _add_strong_points_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add strong points section."""
        doc.add_paragraph("")
        
        strong_title = doc.add_paragraph("💪 STRONG POINTS")
        try:
            strong_title.style = 'CustomHeading'
        except:
            strong_title.runs[0].bold = True
        
        strong_points = analysis_content.get("strong_points", [])
        if not strong_points:
            doc.add_paragraph("No specific strong points identified.")
            return
        
        for i, point in enumerate(strong_points, 1):
            point_para = doc.add_paragraph()
            point_para.add_run(f"{i}. {point.get('point', 'Strong point')}").bold = True
            
            category_para = doc.add_paragraph(f"Category: {point.get('category', 'General')}")
            category_para.runs[0].italic = True
            
            if point.get('explanation'):
                doc.add_paragraph(point['explanation'])
            
            if point.get('leverage'):
                leverage_para = doc.add_paragraph()
                leverage_para.add_run("💡 How to leverage: ").bold = True
                leverage_para.add_run(point['leverage'])
            
            doc.add_paragraph("")

    def _add_weak_points_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add weak points section."""
        doc.add_paragraph("")
        
        weak_title = doc.add_paragraph("⚠️ AREAS FOR IMPROVEMENT")
        try:
            weak_title.style = 'CustomHeading'
        except:
            weak_title.runs[0].bold = True
        
        weak_points = analysis_content.get("weak_points", [])
        if not weak_points:
            doc.add_paragraph("No significant weak points identified.")
            return
        
        for i, point in enumerate(weak_points, 1):
            point_para = doc.add_paragraph()
            point_para.add_run(f"{i}. {point.get('point', 'Improvement area')}").bold = True
            
            category_para = doc.add_paragraph(f"Category: {point.get('category', 'General')}")
            category_para.runs[0].italic = True
            
            if point.get('explanation'):
                doc.add_paragraph(point['explanation'])
            
            if point.get('impact'):
                impact_para = doc.add_paragraph()
                impact_para.add_run("📊 Impact: ").bold = True
                impact_para.add_run(point['impact'])
            
            doc.add_paragraph("")

    def _add_recommendations_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add recommendations section."""
        doc.add_paragraph("")
        
        rec_title = doc.add_paragraph("🚀 SPECIFIC RECOMMENDATIONS")
        try:
            rec_title.style = 'CustomHeading'
        except:
            rec_title.runs[0].bold = True
        
        recommendations = analysis_content.get("recommendations", [])
        if not recommendations:
            doc.add_paragraph("No specific recommendations available.")
            return
        
        for i, rec in enumerate(recommendations, 1):
            priority = rec.get('priority', 'Medium')
            priority_emoji = "🔴" if priority == "High" else "🟡" if priority == "Medium" else "🟢"
            
            rec_para = doc.add_paragraph()
            rec_para.add_run(f"{i}. {rec.get('recommendation', 'Recommendation')} {priority_emoji}").bold = True
            
            category_para = doc.add_paragraph(f"{rec.get('category', 'General')} - {priority} Priority")
            category_para.runs[0].italic = True
            
            if rec.get('action'):
                action_para = doc.add_paragraph()
                action_para.add_run("Action: ").bold = True
                action_para.add_run(rec['action'])
            
            if rec.get('impact'):
                impact_para = doc.add_paragraph()
                impact_para.add_run("Impact: ").bold = True
                impact_para.add_run(rec['impact'])
            
            if rec.get('implementation'):
                impl_para = doc.add_paragraph()
                impl_para.add_run("Implementation: ").bold = True
                impl_para.add_run(rec['implementation'])
            
            doc.add_paragraph("")

    def _add_interview_prep_section(self, doc: Document, analysis_content: Dict[str, Any]) -> None:
        """Add interview preparation section."""
        doc.add_page_break()
        
        prep_title = doc.add_paragraph("📋 INTERVIEW PREPARATION GUIDE")
        try:
            prep_title.style = 'CustomHeading'
        except:
            prep_title.runs[0].bold = True
        
        interview_prep = analysis_content.get("interview_preparation", {})
        
        # Company research
        if "company_research" in interview_prep:
            research = interview_prep["company_research"]
            doc.add_paragraph("Company Research:")
            for key, value in research.items():
                if value:
                    bullet_para = doc.add_paragraph(f"• {value}")
        
        # Common questions
        if "common_questions" in interview_prep:
            doc.add_paragraph("")
            doc.add_paragraph("Common Interview Questions to Prepare:")
            for question in interview_prep["common_questions"]:
                doc.add_paragraph(f"• {question}")
        
        # STAR method
        if "star_method" in interview_prep:
            doc.add_paragraph("")
            star_title = doc.add_paragraph("STAR Method for Behavioral Questions:")
            star_title.runs[0].bold = True
            
            star = interview_prep["star_method"]
            for key, value in star.items():
                para = doc.add_paragraph()
                para.add_run(f"{key.title()}: ").bold = True
                para.add_run(value)
        
        # Questions to ask
        if "questions_to_ask" in interview_prep:
            doc.add_paragraph("")
            doc.add_paragraph("Questions to Ask the Interviewer:")
            for question in interview_prep["questions_to_ask"]:
                doc.add_paragraph(f"• {question}")
        
        # Based on analysis
        if "based_on_analysis" in interview_prep:
            based_on = interview_prep["based_on_analysis"]
            if based_on.get("leverage_strengths"):
                doc.add_paragraph("")
                doc.add_paragraph("Leverage These Strengths:")
                for strength in based_on["leverage_strengths"]:
                    if strength:
                        doc.add_paragraph(f"• {strength}")
            
            if based_on.get("address_concerns"):
                doc.add_paragraph("")
                doc.add_paragraph("Address These Potential Concerns:")
                for concern in based_on["address_concerns"]:
                    if concern:
                        doc.add_paragraph(f"• {concern}")