"""DOCX parser for extracting text from Word documents."""

import logging
from typing import Optional
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


class DocxParser:
    """Parser for extracting text from Word documents."""

    def __init__(self) -> None:
        """Initialize the DOCX parser."""
        self.supported_extensions = [".docx", ".doc"]

    def parse(self, file_path: str) -> Optional[str]:
        """
        Parse DOCX file and extract text content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content or None if parsing fails
        """
        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return None
                
            if path.suffix.lower() not in self.supported_extensions:
                logger.error(f"Unsupported file format: {path.suffix}")
                return None

            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return None
                
            return text.strip()
                
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return None

    def is_supported(self, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        return Path(file_path).suffix.lower() in self.supported_extensions